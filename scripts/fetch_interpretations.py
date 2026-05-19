"""
司法解释批量爬取脚本。

从 flk.npc.gov.cn 下载所有司法解释，保存为 .docx 到 data/司法解释/。

用法：
    python scripts/fetch_interpretations.py [--use-browser]

选项：
    --use-browser    使用 Playwright 浏览器下载全文（需先执行 playwright install chromium）
"""

import argparse
import json
import os
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional

import requests

# 项目根目录
ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT / "data" / "司法解释"

# 司法解释子分类 ID
CATEGORY_CODES = {
    "320": "最高人民法院司法解释",
    "330": "最高人民检察院司法解释",
    "340": "联合发布司法解释",
    "350": "修改、废止的决定",
}

API_BASE = "https://flk.npc.gov.cn"

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "application/json, text/plain, */*",
    "Content-Type": "application/json",
    "Origin": API_BASE,
    "Referer": f"{API_BASE}/",
}

PAGE_SIZE = 20  # 每页最大条数（服务器限制）
REQUEST_DELAY = 0.5  # 请求间隔（秒）
MAX_RETRIES = 3  # 最大重试次数


def safe_filename(title: str, max_len: int = 80) -> str:
    """将标题转为安全文件名。"""
    # 替换 Windows 不允许的字符
    for ch in r'\/:*?"<>|':
        title = title.replace(ch, "_")
    # 去除首尾空白和点
    title = title.strip(". ")
    if len(title) > max_len:
        title = title[:max_len]
    return title


def fetch_list(session: requests.Session, code_id: str, page: int) -> dict:
    """获取一页列表数据。"""
    payload = {
        "searchRange": 1,
        "searchType": 2,
        "searchContent": "",
        "flfgCodeId": [code_id],
        "page": page,
        "size": PAGE_SIZE,
    }
    for attempt in range(MAX_RETRIES):
        try:
            r = session.post(f"{API_BASE}/law-search/search/list", json=payload, timeout=15)
            r.raise_for_status()
            return r.json()
        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                time.sleep(REQUEST_DELAY * (attempt + 1))
            else:
                print(f"  [ERROR] 列表请求失败 (page={page}): {e}")
                return {"total": 0, "rows": []}


def fetch_detail(session: requests.Session, bbbs: str) -> dict:
    """获取单条详情。"""
    for attempt in range(MAX_RETRIES):
        try:
            r = session.get(
                f"{API_BASE}/law-search/search/flfgDetails",
                params={"bbbs": bbbs},
                timeout=15,
            )
            r.raise_for_status()
            data = r.json()
            return data.get("data", {})
        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                time.sleep(REQUEST_DELAY * (attempt + 1))
            else:
                print(f"  [ERROR] 详情请求失败 (bbbs={bbbs}): {e}")
                return {}


def fetch_full_text_playwright(page, bbbs: str) -> Optional[str]:
    """使用 Playwright 从详情页提取全文。"""
    try:
        url = f"{API_BASE}/detail?bbbs={bbbs}"
        page.goto(url, wait_until="networkidle", timeout=30000)
        time.sleep(2)

        # 尝试从 reader iframe 或页面内容中提取文本
        # 1. 尝试获取 reader 内容
        text = page.evaluate("""() => {
            // 尝试获取页面中的文本内容
            const contentEl = document.querySelector('.content-panel') ||
                              document.querySelector('.detail-content') ||
                              document.querySelector('[class*="content"]') ||
                              document.querySelector('.article-body');
            if (contentEl) return contentEl.innerText;

            // 尝试获取所有段落
            const paragraphs = document.querySelectorAll('p');
            if (paragraphs.length > 0) {
                return Array.from(paragraphs).map(p => p.innerText).join('\\n');
            }

            // 尝试从 iframe 获取
            const iframes = document.querySelectorAll('iframe');
            for (const iframe of iframes) {
                try {
                    const doc = iframe.contentDocument;
                    if (doc) return doc.body.innerText;
                } catch(e) {}
            }

            return null;
        }""")
        return text
    except Exception as e:
        print(f"  [WARN] Playwright 提取失败: {e}")
        return None


def save_as_docx(filepath: Path, title: str, category: str, detail: dict, full_text: Optional[str] = None):
    """将司法解释保存为 .docx 文件。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # 标题
    doc.add_heading(title, level=1)

    # 元数据段
    meta_items = []
    if detail.get("zdjgName"):
        meta_items.append(f"发布机关：{detail['zdjgName']}")
    if detail.get("gbrq"):
        meta_items.append(f"颁布日期：{detail['gbrq']}")
    if detail.get("sxrq"):
        meta_items.append(f"施行日期：{detail['sxrq']}")
    if detail.get("flxz"):
        meta_items.append(f"法律性质：{detail['flxz']}")
    if category:
        meta_items.append(f"分类：{category}")

    if meta_items:
        meta_para = doc.add_paragraph()
        meta_para.add_run("\n".join(meta_items)).font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # 正文
    if full_text:
        for line in full_text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("（暂无全文内容，仅保留元数据。完整文本请访问 https://flk.npc.gov.cn 查看。）")

    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))


def save_metadata_json(filepath: Path, items: List[dict]):
    """保存元数据 JSON（用于调试和追踪）。"""
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=2)


def get_existing_titles(data_dir: Path) -> set:
    """获取已下载的文件标题（用于断点续传）。"""
    existing = set()
    if data_dir.exists():
        for f in data_dir.glob("*.docx"):
            # 文件名格式：标题_日期.docx，去掉日期部分
            name = f.stem
            # 去掉末尾的 _YYYYMMDD
            if len(name) > 9 and name[-9] == "_" and name[-8:].isdigit():
                name = name[:-9]
            existing.add(name)
    return existing


def main():
    parser = argparse.ArgumentParser(description="下载司法解释")
    parser.add_argument("--use-browser", action="store_true", help="使用 Playwright 下载全文")
    parser.add_argument("--limit", type=int, default=0, help="每类最多下载条数（0=不限制）")
    parser.add_argument("--output", type=str, default=str(DATA_DIR), help="输出目录")
    args = parser.parse_args()

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 元数据 JSON 路径
    meta_path = output_dir / "_metadata.json"

    # Playwright 初始化（如果需要）
    browser = None
    page = None
    if args.use_browser:
        try:
            from playwright.sync_api import sync_playwright
            pw = sync_playwright().start()
            browser = pw.chromium.launch(headless=True)
            page = browser.new_page()
            print("Playwright 浏览器已启动")
        except Exception as e:
            print(f"Playwright 初始化失败: {e}")
            print("将仅保存元数据，不获取全文")
            args.use_browser = False

    session = requests.Session()
    session.headers.update(HEADERS)

    existing = get_existing_titles(output_dir)

    # 加载已有元数据（累积模式）
    all_metadata = []
    if meta_path.exists():
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                all_metadata = json.load(f)
            print(f"已加载 {len(all_metadata)} 条已有元数据")
        except Exception:
            pass
    existing_bbbs = {m["bbbs"] for m in all_metadata}

    stats = {"total": 0, "downloaded": 0, "skipped": 0, "failed": 0}

    for code_id, category_name in CATEGORY_CODES.items():
        print(f"\n{'='*60}")
        print(f"分类: {category_name} (codeId={code_id})")
        print(f"{'='*60}")

        # 获取第一页，拿到总数
        first_page = fetch_list(session, code_id, 1)
        total = first_page.get("total", 0)
        rows = first_page.get("rows", [])
        print(f"  总条数: {total}")

        if total == 0:
            continue

        # 收集所有条目
        all_rows = list(rows)
        total_pages = (total + PAGE_SIZE - 1) // PAGE_SIZE

        for page_num in range(2, total_pages + 1):
            time.sleep(REQUEST_DELAY)
            page_data = fetch_list(session, code_id, page_num)
            page_rows = page_data.get("rows", [])
            all_rows.extend(page_rows)
            if page_num % 10 == 0:
                print(f"  已获取 {len(all_rows)}/{total} 条...")

        print(f"  共获取 {len(all_rows)} 条")

        if args.limit > 0:
            all_rows = all_rows[:args.limit]
            print(f"  限制为 {args.limit} 条")

        # 逐条处理
        for i, row in enumerate(all_rows):
            stats["total"] += 1
            title = row.get("title", "未知标题")
            bbbs = row.get("bbbs", "")
            gbrq = row.get("gbrq", "")
            sxrq = row.get("sxrq", "")

            # 按 bbbs 去重（API 分页可能返回重复条目）
            if bbbs in existing_bbbs:
                stats["skipped"] += 1
                continue

            # 文件名
            date_str = sxrq.replace("-", "") if sxrq else gbrq.replace("-", "") if gbrq else "unknown"
            filename = safe_filename(f"{title}_{date_str}")
            filepath = output_dir / f"{filename}.docx"

            # 断点续传
            safe_title = safe_filename(title)
            if safe_title in existing and filepath.exists():
                stats["skipped"] += 1
                # 补充缺失的元数据
                all_metadata.append({
                    "bbbs": bbbs,
                    "title": title,
                    "category": category_name,
                    "code_id": code_id,
                    "gbrq": gbrq,
                    "sxrq": sxrq,
                    "zdjgName": "",
                    "flxz": "",
                    "oss_word_path": "",
                    "filename": f"{filename}.docx",
                    "has_full_text": False,
                })
                existing_bbbs.add(bbbs)
                continue

            print(f"  [{i+1}/{len(all_rows)}] {title[:50]}...", end="")

            # 获取详情
            time.sleep(REQUEST_DELAY)
            detail = fetch_detail(session, bbbs)
            if not detail:
                print(" [FAIL: 无详情]")
                stats["failed"] += 1
                continue

            # 尝试获取全文
            full_text = None
            if args.use_browser and page:
                full_text = fetch_full_text_playwright(page, bbbs)

            # 保存 .docx
            try:
                save_as_docx(filepath, title, category_name, detail, full_text)
                stats["downloaded"] += 1
                print(" [OK]")
            except Exception as e:
                print(f" [FAIL: {e}]")
                stats["failed"] += 1
                continue

            # 收集元数据
            all_metadata.append({
                "bbbs": bbbs,
                "title": title,
                "category": category_name,
                "code_id": code_id,
                "gbrq": gbrq,
                "sxrq": sxrq,
                "zdjgName": detail.get("zdjgName", ""),
                "flxz": detail.get("flxz", ""),
                "oss_word_path": detail.get("ossFile", {}).get("ossWordPath", ""),
                "filename": f"{filename}.docx",
                "has_full_text": full_text is not None,
            })
            existing_bbbs.add(bbbs)

    # 保存元数据
    if all_metadata:
        save_metadata_json(meta_path, all_metadata)

    # 关闭浏览器
    if browser:
        browser.close()

    # 统计
    print(f"\n{'='*60}")
    print(f"完成！统计:")
    print(f"  总条数: {stats['total']}")
    print(f"  已下载: {stats['downloaded']}")
    print(f"  已跳过: {stats['skipped']}")
    print(f"  失败:   {stats['failed']}")
    print(f"  输出目录: {output_dir}")
    if all_metadata:
        print(f"  元数据: {meta_path}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
